"""
Alternative vector store implementation for environments with SQLite compatibility issues.
This module provides a fallback FAISS-based vector store when ChromaDB is not available.
"""

import os
import pickle
from typing import List, Dict, Any, Optional
from pathlib import Path
import numpy as np
from loguru import logger

# Vector Store Dependencies
try:
    from langchain_community.vectorstores import FAISS
    FAISS_AVAILABLE = True
except ImportError:
    try:
        # Try alternative import path
        from langchain.vectorstores import FAISS
        FAISS_AVAILABLE = True
    except ImportError:
        FAISS_AVAILABLE = False
        logger.warning("FAISS not available, using in-memory vector store")

try:
    from langchain.schema import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError as e:
    logger.error(f"Critical langchain imports failed: {e}")
    raise

from src.api_providers import APIProviderFactory


class AlternativeVectorStore:
    """
    Alternative vector store implementation using FAISS or in-memory storage.
    Used as fallback when ChromaDB has SQLite compatibility issues.
    """
    
    def __init__(self, persist_directory: str = "data/faiss_db"):
        self.persist_directory = Path(persist_directory)
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        
        self.embeddings = APIProviderFactory.get_embeddings()
        self.vectorstore = None
        self.documents = []
        
        # Try to load existing vector store
        self._load_vectorstore()
    
    def ingest_file(self, file_path: str) -> Dict[str, Any]:
        """
        Ingest a single file into the alternative vector store.
        
        Args:
            file_path: Path to the file to ingest
            
        Returns:
            Ingestion result summary
        """
        try:
            # Process document using the document processor (which handles file reading)
            try:
                from src.ingestion.document_processor import process_document
                processed_doc = process_document(file_path)
            except ImportError:
                # Fallback to simple processing if document processor not available
                content = self._read_file_content(file_path)
                processed_doc = {
                    'document_id': f"doc_{hash(content)%10000}",
                    'sections': [{'content': content, 'section_number': '1'}],
                    'document_type': 'document'
                }
            
            # Create Document objects
            documents = []
            for section in processed_doc.get('sections', []):
                doc = Document(
                    page_content=section['content'],
                    metadata={
                        'source': file_path,
                        'section_number': section['section_number'],
                        'document_type': processed_doc.get('document_type', 'unknown'),
                        'filename': Path(file_path).name
                    }
                )
                documents.append(doc)
            
            # Add to vector store
            doc_ids = self.add_documents(documents)
            
            return {
                'success': True,
                'document_id': processed_doc.get('document_id'),
                'sections_count': len(processed_doc.get('sections', [])),
                'chunks_count': len(documents),
                'vector_store_type': 'Alternative'
            }
            
        except Exception as e:
            logger.error(f"Failed to ingest file {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'document_id': None
            }
    
    def _read_file_content(self, file_path: str) -> str:
        """Read content from file based on file type."""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            try:
                import PyMuPDF as fitz
                doc = fitz.open(str(file_path))
                content = ""
                for page in doc:
                    content += page.get_text()
                doc.close()
                return content
            except ImportError:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(str(file_path))
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text()
                    return content
                except ImportError:
                    raise ImportError("No PDF reader available. Install PyMuPDF or pypdf.")
        
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                return content
            except ImportError:
                raise ImportError("python-docx not available for Word documents.")
        
        elif file_path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

    def add_documents(self, documents: List[Document]) -> List[str]:
        """Add documents to the vector store."""
        try:
            if FAISS_AVAILABLE:
                if self.vectorstore is None:
                    # Create new FAISS vector store
                    self.vectorstore = FAISS.from_documents(documents, self.embeddings)
                else:
                    # Add to existing vector store
                    self.vectorstore.add_documents(documents)
                
                # Save the vector store
                self._save_vectorstore()
                
            else:
                # Fallback to in-memory storage
                self.documents.extend(documents)
                self._save_documents()
            
            return [f"doc_{i}" for i in range(len(documents))]
            
        except Exception as e:
            logger.error(f"Error adding documents to alternative vector store: {e}")
            raise
    
    def similarity_search(self, query: str, k: int = 4) -> List[Document]:
        """Search for similar documents."""
        try:
            if FAISS_AVAILABLE and self.vectorstore is not None:
                return self.vectorstore.similarity_search(query, k=k)
            else:
                # Fallback: simple text matching
                return self._simple_text_search(query, k=k)
                
        except Exception as e:
            logger.error(f"Error in similarity search: {e}")
            return []
    
    def _simple_text_search(self, query: str, k: int = 4) -> List[Document]:
        """Simple text-based search fallback."""
        query_lower = query.lower()
        scored_docs = []
        
        for doc in self.documents:
            content_lower = doc.page_content.lower()
            score = 0
            
            # Simple scoring based on keyword matches
            query_words = query_lower.split()
            for word in query_words:
                score += content_lower.count(word)
            
            if score > 0:
                scored_docs.append((score, doc))
        
        # Sort by score and return top k
        scored_docs.sort(key=lambda x: x[0], reverse=True)
        return [doc for _, doc in scored_docs[:k]]
    
    def _save_vectorstore(self):
        """Save FAISS vector store to disk."""
        if FAISS_AVAILABLE and self.vectorstore is not None:
            try:
                self.vectorstore.save_local(str(self.persist_directory))
                logger.info("Vector store saved successfully")
            except Exception as e:
                logger.error(f"Error saving vector store: {e}")
    
    def _load_vectorstore(self):
        """Load FAISS vector store from disk."""
        if FAISS_AVAILABLE:
            try:
                index_path = self.persist_directory / "index.faiss"
                if index_path.exists():
                    self.vectorstore = FAISS.load_local(
                        str(self.persist_directory), 
                        self.embeddings,
                        allow_dangerous_deserialization=True
                    )
                    logger.info("Vector store loaded successfully")
            except Exception as e:
                logger.warning(f"Could not load existing vector store: {e}")
    
    def _save_documents(self):
        """Save documents to pickle file as fallback."""
        try:
            docs_path = self.persist_directory / "documents.pkl"
            with open(docs_path, 'wb') as f:
                pickle.dump(self.documents, f)
        except Exception as e:
            logger.error(f"Error saving documents: {e}")
    
    def _load_documents(self):
        """Load documents from pickle file."""
        try:
            docs_path = self.persist_directory / "documents.pkl"
            if docs_path.exists():
                with open(docs_path, 'rb') as f:
                    self.documents = pickle.load(f)
                logger.info(f"Loaded {len(self.documents)} documents from fallback storage")
        except Exception as e:
            logger.warning(f"Could not load documents from fallback storage: {e}")
    
    def get_collection_info(self) -> Dict[str, Any]:
        """Get information about the vector store collection."""
        if FAISS_AVAILABLE and self.vectorstore is not None:
            return {
                'type': 'FAISS',
                'document_count': len(self.documents),
                'index_size': self.vectorstore.index.ntotal if hasattr(self.vectorstore.index, 'ntotal') else 'Unknown'
            }
        else:
            return {
                'type': 'In-Memory',
                'document_count': len(self.documents),
                'index_size': 'N/A'
            }


def create_alternative_vector_store(persist_directory: str = "data/faiss_db") -> AlternativeVectorStore:
    """Create and return an alternative vector store instance."""
    return AlternativeVectorStore(persist_directory)
