"""
Compatibility layer for document ingestion across different vector store implementations.
Provides a unified interface regardless of which vector store is available.
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
from loguru import logger

try:
    from langchain.schema import Document
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError as e:
    logger.error(f"Critical langchain dependencies missing: {e}")
    raise

from src.utils import DocumentUtils
from src.ingestion.document_processor import process_document, process_documents_batch
from src.api_providers import APIProviderFactory

# Try to determine which vector store implementation to use
VECTOR_STORE_TYPE = None
vector_store_class = None

try:
    from src.ingestion.vector_store import VectorStoreManager, CHROMA_AVAILABLE
    if CHROMA_AVAILABLE:
        VECTOR_STORE_TYPE = "ChromaDB"
        vector_store_class = VectorStoreManager
    else:
        raise ImportError("ChromaDB not available")
except ImportError:
    try:
        from src.ingestion.alternative_vector_store import create_alternative_vector_store
        VECTOR_STORE_TYPE = "Alternative"
        vector_store_class = create_alternative_vector_store
        logger.info("Using alternative vector store implementation")
    except ImportError as e:
        logger.error(f"No vector store implementation available: {e}")
        raise


class UnifiedDocumentIngestionPipeline:
    """
    Unified document ingestion pipeline that works with any available vector store.
    """
    
    def __init__(self, persist_directory: Optional[str] = None):
        """Initialize the ingestion pipeline."""
        self.persist_directory = persist_directory or "data/vector_db"
        self.vector_store = None
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Initialize vector store based on available implementation
        self._initialize_vector_store()
    
    def _initialize_vector_store(self):
        """Initialize the appropriate vector store implementation."""
        try:
            if VECTOR_STORE_TYPE == "ChromaDB":
                self.vector_store = vector_store_class(self.persist_directory)
                logger.info("Initialized ChromaDB vector store")
            elif VECTOR_STORE_TYPE == "Alternative":
                self.vector_store = vector_store_class(self.persist_directory)
                logger.info("Initialized alternative vector store")
            else:
                raise RuntimeError("No vector store implementation available")
        except Exception as e:
            logger.error(f"Failed to initialize vector store: {e}")
            raise
    
    def ingest_file(self, file_path: str) -> Dict[str, Any]:
        """
        Ingest a file - reads content and processes it.
        Maintains compatibility with existing code that only passes file_path.
        """
        try:
            # Read file content based on file type
            content = self._read_file_content(file_path)
            return self.process_and_ingest_document(file_path, content)
        except Exception as e:
            logger.error(f"Failed to ingest file {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'document_id': None
            }
    
    def _read_file_content(self, file_path: str) -> str:
        """Read content from file based on file type."""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == '.pdf':
            try:
                import PyMuPDF as fitz
                doc = fitz.open(str(file_path))
                content = ""
                for page in doc:
                    content += page.get_text()
                doc.close()
                return content
            except ImportError:
                try:
                    from pypdf import PdfReader
                    reader = PdfReader(str(file_path))
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text()
                    return content
                except ImportError:
                    raise ImportError("No PDF reader available. Install PyMuPDF or pypdf.")
        
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                return content
            except ImportError:
                raise ImportError("python-docx not available for Word documents.")
        
        elif file_path.suffix.lower() == '.txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def ingest_file_with_content(self, file_path: str, content: str) -> Dict[str, Any]:
        """
        Ingest a file with pre-read content.
        Use this method when you already have the file content.
        """
        return self.process_and_ingest_document(file_path, content)
    
    def ingest_documents(self, documents: List[Document]) -> List[str]:
        """
        Ingest pre-processed documents directly.
        """
        try:
            # Add to vector store
            if hasattr(self.vector_store, 'add_documents'):
                doc_ids = self.vector_store.add_documents(documents)
            elif hasattr(self.vector_store, 'ingest_documents'):
                doc_ids = self.vector_store.ingest_documents(documents)
            else:
                raise AttributeError("Vector store has no document ingestion method")
            
            return doc_ids
            
        except Exception as e:
            logger.error(f"Failed to ingest documents: {e}")
            raise

    def process_and_ingest_document(self, file_path: str, content: str) -> Dict[str, Any]:
        """Process and ingest a single document."""
        try:
            # Process document
            processed_doc = process_document(file_path, content)
            
            # Create Document objects for vector store
            documents = []
            for section in processed_doc.get('sections', []):
                doc = Document(
                    page_content=section['content'],
                    metadata={
                        'source': file_path,
                        'section_number': section['section_number'],
                        'document_type': processed_doc.get('document_type', 'unknown'),
                        'filename': Path(file_path).name
                    }
                )
                documents.append(doc)
            
            # Split documents if needed
            split_docs = self.text_splitter.split_documents(documents)
            
            # Add to vector store
            if hasattr(self.vector_store, 'add_documents'):
                doc_ids = self.vector_store.add_documents(split_docs)
            elif hasattr(self.vector_store, 'ingest_documents'):
                doc_ids = self.vector_store.ingest_documents(split_docs)
            else:
                raise AttributeError("Vector store has no document ingestion method")
            
            return {
                'success': True,
                'document_id': processed_doc.get('document_id'),
                'sections_count': len(processed_doc.get('sections', [])),
                'chunks_count': len(split_docs),
                'vector_store_type': VECTOR_STORE_TYPE
            }
            
        except Exception as e:
            logger.error(f"Failed to process and ingest document {file_path}: {e}")
            return {
                'success': False,
                'error': str(e),
                'document_id': None
            }
    
    def search_documents(self, query: str, k: int = 4) -> List[Document]:
        """Search for relevant documents."""
        try:
            if hasattr(self.vector_store, 'similarity_search'):
                return self.vector_store.similarity_search(query, k=k)
            elif hasattr(self.vector_store, 'search'):
                return self.vector_store.search(query, k=k)
            else:
                logger.warning("Vector store has no search method")
                return []
        except Exception as e:
            logger.error(f"Search failed: {e}")
            return []
    
    def get_collection_info(self) -> Dict[str, Any]:
        """Get information about the vector store collection."""
        try:
            if hasattr(self.vector_store, 'get_collection_info'):
                info = self.vector_store.get_collection_info()
                info['vector_store_type'] = VECTOR_STORE_TYPE
                return info
            else:
                return {
                    'vector_store_type': VECTOR_STORE_TYPE,
                    'status': 'available',
                    'document_count': 'unknown'
                }
        except Exception as e:
            logger.error(f"Failed to get collection info: {e}")
            return {
                'vector_store_type': VECTOR_STORE_TYPE,
                'status': 'error',
                'error': str(e)
            }


# Maintain backward compatibility
DocumentIngestionPipeline = UnifiedDocumentIngestionPipeline
