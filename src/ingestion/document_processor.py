"""
Document processors for different file formats.
Handles PDF, DOCX, and TXT files with legal document-specific processing.
"""

import io
from pathlib import Path
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod

try:
    import fitz  # PyMuPDF
    from docx import Document
    from loguru import logger
except ImportError as e:
    logger = None
    print(f"Import warning: {e}")

from src.utils import DocumentUtils
from config.settings import settings


class BaseDocumentProcessor(ABC):
    """Abstract base class for document processors."""
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text from document."""
        pass
    
    @abstractmethod
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from document."""
        pass


class PDFProcessor(BaseDocumentProcessor):
    """Processor for PDF documents using PyMuPDF."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(file_path)
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                
                # Add page break marker
                text_content.append(f"\n--- Page {page_num + 1} ---\n{text}")
            
            doc.close()
            return "\n".join(text_content)
            
        except Exception as e:
            if logger:
                logger.error(f"Error extracting text from PDF {file_path}: {e}")
            raise RuntimeError(f"Failed to extract text from PDF: {e}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file."""
        try:
            doc = fitz.open(file_path)
            metadata = doc.metadata
            
            result = {
                'title': metadata.get('title', ''),
                'author': metadata.get('author', ''),
                'subject': metadata.get('subject', ''),
                'creator': metadata.get('creator', ''),
                'producer': metadata.get('producer', ''),
                'creation_date': metadata.get('creationDate', ''),
                'modification_date': metadata.get('modDate', ''),
                'page_count': len(doc),
                'file_size': Path(file_path).stat().st_size,
                'file_type': 'PDF'
            }
            
            doc.close()
            return result
            
        except Exception as e:
            if logger:
                logger.error(f"Error extracting metadata from PDF {file_path}: {e}")
            return {'file_type': 'PDF', 'error': str(e)}


class DOCXProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents using python-docx."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            if logger:
                logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise RuntimeError(f"Failed to extract text from DOCX: {e}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file."""
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            result = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'creator': core_props.author or '',
                'creation_date': str(core_props.created) if core_props.created else '',
                'modification_date': str(core_props.modified) if core_props.modified else '',
                'revision': core_props.revision or '',
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'file_size': Path(file_path).stat().st_size,
                'file_type': 'DOCX'
            }
            
            return result
            
        except Exception as e:
            if logger:
                logger.error(f"Error extracting metadata from DOCX {file_path}: {e}")
            return {'file_type': 'DOCX', 'error': str(e)}


class TXTProcessor(BaseDocumentProcessor):
    """Processor for plain text documents."""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                if logger:
                    logger.error(f"Error reading TXT file {file_path}: {e}")
                raise RuntimeError(f"Failed to read TXT file: {e}")
        except Exception as e:
            if logger:
                logger.error(f"Error extracting text from TXT {file_path}: {e}")
            raise RuntimeError(f"Failed to extract text from TXT: {e}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from TXT file."""
        try:
            file_path_obj = Path(file_path)
            stat = file_path_obj.stat()
            
            result = {
                'title': file_path_obj.stem,
                'creation_date': str(stat.st_ctime),
                'modification_date': str(stat.st_mtime),
                'file_size': stat.st_size,
                'file_type': 'TXT'
            }
            
            return result
            
        except Exception as e:
            if logger:
                logger.error(f"Error extracting metadata from TXT {file_path}: {e}")
            return {'file_type': 'TXT', 'error': str(e)}


class DocumentProcessorFactory:
    """Factory class to get appropriate document processor."""
    
    _processors = {
        '.pdf': PDFProcessor,
        '.docx': DOCXProcessor,
        '.doc': DOCXProcessor,  # Treat .doc as .docx
        '.txt': TXTProcessor,
    }
    
    @classmethod
    def get_processor(cls, file_path: str) -> BaseDocumentProcessor:
        """Get appropriate processor for file type."""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in cls._processors:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        processor_class = cls._processors[file_extension]
        return processor_class()
    
    @classmethod
    def supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions."""
        return list(cls._processors.keys())


def process_document(file_path: str) -> Dict[str, Any]:
    """
    Process a document and extract text, metadata, and sections.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Dictionary containing processed document data
    """
    try:
        processor = DocumentProcessorFactory.get_processor(file_path)
        
        # Extract text and metadata
        text = processor.extract_text(file_path)
        metadata = processor.extract_metadata(file_path)
        
        # Clean and process text
        cleaned_text = DocumentUtils.clean_text(text)
        
        # Extract sections with error handling
        try:
            sections = DocumentUtils.extract_sections(cleaned_text)
        except Exception as section_error:
            if logger:
                logger.warning(f"Error extracting sections from {file_path}: {section_error}. Using simple chunking.")
            # Fallback to simple chunking
            chunks = DocumentUtils.split_into_chunks(cleaned_text, settings.chunk_size)
            sections = [
                {
                    'section_number': str(i + 1),
                    'content': chunk,
                    'type': 'chunk'
                }
                for i, chunk in enumerate(chunks)
            ]
        
        # Generate document ID
        doc_id = DocumentUtils.generate_document_id(cleaned_text, Path(file_path).name)
        
        result = {
            'document_id': doc_id,
            'filename': Path(file_path).name,
            'filepath': file_path,
            'text': cleaned_text,
            'sections': sections,
            'metadata': metadata,
            'token_count': DocumentUtils.count_tokens(cleaned_text),
            'section_count': len(sections)
        }
        
        if logger:
            logger.info(f"Successfully processed document: {file_path} "
                       f"({len(sections)} sections, {result['token_count']} tokens)")
        
        return result
        
    except Exception as e:
        if logger:
            logger.error(f"Failed to process document {file_path}: {e}")
        raise


def process_documents_batch(file_paths: List[str]) -> List[Dict[str, Any]]:
    """
    Process multiple documents in batch.
    
    Args:
        file_paths: List of file paths to process
        
    Returns:
        List of processed document data dictionaries
    """
    results = []
    
    for file_path in file_paths:
        try:
            result = process_document(file_path)
            results.append(result)
        except Exception as e:
            if logger:
                logger.error(f"Failed to process document {file_path}: {e}")
            # Continue processing other documents
            continue
    
    return results
